"""08 个人文档：文本解析（纯文本 / HTML / PDF / Word，二进制显式拒绝、不抛 500）。

约定（技术文档 §5.3）：
- ``.txt`` / ``.md`` / ``.markdown`` / ``.html`` / ``.htm`` → 解析为纯文本；
- ``.pdf`` → 用 ``pypdf`` 抽取每页文本（依赖 ``pypdf``，缺失时给出明确提示）；
- ``.docx`` → 用 ``python-docx`` 抽取段落与表格文本（OOXML）；
- ``.doc`` / ``.docs``（老版 Word OLE2 二进制）→ 无需外部依赖，从字节中尽力抽取
  正文（UTF-16LE 中文串 + ASCII 英文串）；扫描件/纯图片无文本层时抛
  :class:`EmptyDocument` 交由调用方降级 ``status=error``；
- 其它扩展名（``.xlsx`` …）→ 抛 :class:`UnsupportedDocument`，
  调用方据此把该文档标记为 ``status=error`` 并给出明确文案，**不影响同批次其它文件**。

HTML 去标签逻辑与 06 ``store.ingest_document`` 保持一致（同一套正则），
避免两处解析行为漂移。
"""
from __future__ import annotations

import io
import os
import re

SUPPORTED_EXTS = (".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx", ".doc", ".docs")

# 启发式抽取时保留的"文字类"码位之外的常见中英文标点/符号
_PUNCT = set(" ·…—、，。：；！？（）《》“”‘’.,;:!?()'\"+/-")


class UnsupportedDocument(Exception):
    """文档类型不受支持（调用方转 status=error，非 HTTP 500）。"""


class EmptyDocument(Exception):
    """文档解析后无有效文本（空文件 / 全空白）。"""


def is_supported(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in SUPPORTED_EXTS


def strip_html(text: str) -> str:
    text = re.sub(r"<script[\s\S]*?</script>", " ", text)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;", " ", text)
    return re.sub(r"\s+", " ", text)


def parse_bytes(filename: str, raw: bytes) -> str:
    """把上传文件字节解析为纯文本。

    Raises:
        UnsupportedDocument: 扩展名不在 :data:`SUPPORTED_EXTS` 内（或 PDF 依赖缺失）。
        EmptyDocument: 解析结果为空（空文件 / 全空白 / 扫描件无文本层）。
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTS:
        raise UnsupportedDocument(
            f"暂不支持的文档类型：{ext or '(无扩展名)'}"
            f"（本期支持 {'/'.join(SUPPORTED_EXTS)}）"
        )
    if ext == ".pdf":
        return _parse_pdf(raw)
    if ext in (".doc", ".docs", ".docx"):
        return _parse_doc(raw)
    text = raw.decode("utf-8", errors="ignore")
    if ext in (".html", ".htm"):
        text = strip_html(text)
    if not text.strip():
        raise EmptyDocument("文档内容为空，无可解析文本。")
    return text


def _parse_pdf(raw: bytes) -> str:
    """用 ``pypdf`` 抽取每页文本；缺库或扫描件时给明确报错。

    单页抽取异常不阻断整本；全本无文本层（图片型/扫描件 PDF）抛
    :class:`EmptyDocument` 交由调用方降级为 ``status=error``。
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # 依赖缺失提示，而非 500
        raise UnsupportedDocument("PDF 解析依赖 pypdf 未安装（pip install pypdf）") from exc
    import io

    reader = PdfReader(io.BytesIO(raw))
    if not reader.pages:
        raise EmptyDocument("PDF 无页面内容。")
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")  # 单页失败不阻断整本
    text = "\n".join(p.strip() for p in parts).strip()
    if not text:
        raise EmptyDocument("PDF 未提取到可解析文本（可能是扫描件/图片型 PDF）。")
    return text


def _parse_doc(raw: bytes) -> str:
    """解析 Word 文档：``.docx`` 走 python-docx；真实 ``.doc``/``.docs`` 二进制走启发式。

    - ``.docx``（OOXML）：``python-docx`` 直接拿到干净段落与表格文本；
    - ``.doc`` / ``.docs``（OLE2 二进制）：``python-docx`` 会解析失败，落到
      :func:`_extract_doc_text` 从字节里尽力抽取正文（无需 antiword/catdoc/libreoffice）。
    两者都抽不到文本（扫描件/图片型）抛 :class:`EmptyDocument`。
    """
    # 1) 先试 OOXML（.docx，或误用扩展名的 docx 文件）
    try:
        from docx import Document

        try:
            doc = Document(io.BytesIO(raw))
            parts: list[str] = [
                p.text for p in doc.paragraphs if p.text and p.text.strip()
            ]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text)
            text = "\n".join(parts).strip()
            if text:
                return text
        except Exception:
            pass  # 真实 .doc 二进制 → 落到启发式
    except ImportError:
        pass

    # 2) 启发式：从 OLE2 二进制字节里尽力抽取正文
    text = _extract_doc_text(raw)
    if not text:
        raise EmptyDocument(
            "Word 文档未提取到可解析文本（可能是扫描件/图片型 .doc）。"
        )
    return text


def _extract_doc_text(raw: bytes) -> str:
    """老版 ``.doc``/``.docs`` 二进制：从 UTF-16LE 正文流保留文字类码位、丢弃二进制噪声。

    Word 正文主要以 UTF-16LE 存储（中英文混排），直接按该编码解码后保留
    CJK / ASCII 字母数字 / 空白 / 常见中英文标点，丢弃控制符与二进制字节，
    即可还原绝大多数可见文本（无外部依赖，尽力而为）。
    """
    try:
        u16 = raw.decode("utf-16-le", errors="ignore")
    except Exception:
        return ""
    kept: list[str] = []
    for ch in u16:
        o = ord(ch)
        if (
            (0x3400 <= o <= 0x4DBF)        # 扩展 A 汉字
            or (0x4E00 <= o <= 0x9FFF)      # 常用汉字
            or (0x3000 <= o <= 0x303F)      # CJK 标点
            or (0xFF00 <= o <= 0xFFEF)      # 全角字符
            or (0x2010 <= o <= 0x2027)      # 连字符/破折号等
            or (0x2018 <= o <= 0x201F)      # 智能引号
            or ("a" <= ch <= "z")
            or ("A" <= ch <= "Z")
            or ("0" <= ch <= "9")
            or ch in _PUNCT
        ):
            kept.append(ch)
        # 其余（含 \\x00 与二进制）丢弃
    return re.sub(r"\s+", " ", "".join(kept)).strip()


def make_summary(text: str, limit: int = 120) -> str:
    """取正文首段压缩为一行摘要，供前端「基于此文档提问」展示。"""
    flat = re.sub(r"\s+", " ", text).strip()
    return flat[:limit] + ("…" if len(flat) > limit else "")
